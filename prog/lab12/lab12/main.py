import tkinter as tk
from tkinter import messagebox, filedialog
from geometry.rectangle import Rectangle
from geometry.triangle import Triangle
from geometry.trapezoid import Trapezoid
import openpyxl
from docx import Document

class GeometryApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Geometry Calculator")

        # Выбор фигуры
        tk.Label(root, text="Выберите фигуру:").grid(row=0, column=0, padx=10, pady=10)
        self.figure_var = tk.StringVar(value="rectangle")
        tk.Radiobutton(root, text="Прямоугольник", variable=self.figure_var, value="rectangle").grid(row=1, column=0)
        tk.Radiobutton(root, text="Треугольник", variable=self.figure_var, value="triangle").grid(row=2, column=0)
        tk.Radiobutton(root, text="Трапеция", variable=self.figure_var, value="trapezoid").grid(row=3, column=0)

        # Ввод параметров
        self.param_labels = ["Параметр 1", "Параметр 2", "Параметр 3"]
        self.param_entries = [tk.Entry(root) for _ in range(3)]

        for i, entry in enumerate(self.param_entries):
            tk.Label(root, text=self.param_labels[i]).grid(row=i, column=1)
            entry.grid(row=i, column=2)

        # Кнопки
        tk.Button(root, text="Рассчитать", command=self.calculate).grid(row=4, column=0, columnspan=3, pady=10)
        tk.Button(root, text="Сохранить в Excel", command=self.save_to_excel).grid(row=5, column=0, columnspan=3)
        tk.Button(root, text="Сохранить в Word", command=self.save_to_word).grid(row=6, column=0, columnspan=3)

        # Результат
        self.result_text = tk.Text(root, height=5, width=50)
        self.result_text.grid(row=7, column=0, columnspan=3, pady=10)

    def calculate(self):
        try:
            shape = self.figure_var.get()
            params = [float(entry.get()) for entry in self.param_entries if entry.get()]
            if shape == "rectangle":
                rect = Rectangle(*params)
                self.results = {
                    "Площадь": rect.area(),
                    "Радиус описанной окружности": rect.circumradius(),
                    "Радиус вписанной окружности": rect.inradius(),
                }
            elif shape == "triangle":
                tri = Triangle(*params)
                self.results = {
                    "Площадь": tri.area(),
                    "Радиус описанной окружности": tri.circumradius(),
                    "Радиус вписанной окружности": tri.inradius(),
                }
            elif shape == "trapezoid":
                trap = Trapezoid(*params)
                self.results = {
                    "Площадь": trap.area(),
                    "Радиус описанной окружности": trap.circumradius(),
                    "Радиус вписанной окружности": trap.inradius(),
                }
            else:
                raise ValueError("Неверная фигура.")

            self.result_text.delete(1.0, tk.END)
            for key, value in self.results.items():
                self.result_text.insert(tk.END, f"{key}: {value:.2f}\n")
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))

    def save_to_excel(self):
        try:
            filepath = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
            if not filepath:
                return
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.title = "Results"
            for i, (key, value) in enumerate(self.results.items(), start=1):
                sheet[f"A{i}"] = key
                sheet[f"B{i}"] = value
            workbook.save(filepath)
            messagebox.showinfo("Успех", "Результаты сохранены в Excel.")
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))
            
    def save_to_word(self):
        try:
            filepath = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if not filepath:
                return
            doc = Document()
            doc.add_heading("Результаты расчетов", level=1)
            for key, value in self.results.items():
                doc.add_paragraph(f"{key}: {value:.2f}")
            doc.save(filepath)
            messagebox.showinfo("Успех", "Результаты сохранены в Word.")
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))        
if __name__ == "__main__":
    root = tk.Tk()
    app = GeometryApp(root)
    root.mainloop()