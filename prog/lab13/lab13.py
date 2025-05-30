import tkinter as tk
from tkinter import messagebox, filedialog
from abc import ABC, abstractmethod
from docx import Document
import openpyxl
import math


# Абстрактный базовый класс для геометрической фигуры
class Shape(ABC):
    def __init__(self, name):
        self.name = name

    # Абстрактные методы для расчёта площади и радиусов
    @abstractmethod
    def area(self):
        pass

    @abstractmethod
    def inscribed_radius(self):
        pass

    @abstractmethod
    def circumscribed_radius(self):
        pass

    def __str__(self):
        return f"Shape: {self.name}"


# Прямоугольник
class Rectangle(Shape):
    __slots__ = ['length', 'width']

    def __init__(self, length, width):
        super().__init__('Прямоугольник')
        self.length = length
        self.width = width

    def area(self):
        return self.length * self.width

    def inscribed_radius(self):
        return min(self.length, self.width) / 2

    def circumscribed_radius(self):
        return math.sqrt(self.length**2 + self.width**2) / 2

    def __repr__(self):
        return f"Rectangle(length={self.length}, width={self.width})"


# Треугольник
class Triangle(Shape):
    __slots__ = ['a', 'b', 'c']

    def __init__(self, a, b, c):
        super().__init__('Треугольник')
        self.a = a
        self.b = b
        self.c = c

    def area(self):
        s = (self.a + self.b + self.c) / 2
        return math.sqrt(s * (s - self.a) * (s - self.b) * (s - self.c))

    def inscribed_radius(self):
        return self.area() / ((self.a + self.b + self.c) / 2)

    def circumscribed_radius(self):
        return (self.a * self.b * self.c) / (4 * self.area())

    def __repr__(self):
        return f"Triangle(a={self.a}, b={self.b}, c={self.c})"


# Трапеция
class Trapezoid(Shape):
    __slots__ = ['a', 'b', 'h']

    def __init__(self, a, b, h):
        super().__init__('Трапеция')
        self.a = a
        self.b = b
        self.h = h

    def area(self):
        return ((self.a + self.b) / 2) * self.h

    def inscribed_radius(self):
        if self.a == self.b:
            return self.h / 2
        return None  # Нет вписанной окружности для равнобедренной трапеции

    def circumscribed_radius(self):
        return math.sqrt(((self.a + self.b) / 2)**2 + self.h**2)

    def __repr__(self):
        return f"Trapezoid(a={self.a}, b={self.b}, h={self.h})"


# Приложение GUI
class GeometryApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Геометрический калькулятор")

        # Создание переменных
        self.shape = None
        self.result_text = tk.StringVar()

        # Создание интерфейса
        self.create_widgets()

    def create_widgets(self):
        tk.Label(self.root, text="Выберите фигуру:").grid(row=0, column=0, sticky="w")
        self.shape_var = tk.StringVar(value="Прямоугольник")
        tk.OptionMenu(self.root, self.shape_var, "Прямоугольник", "Треугольник", "Трапеция").grid(row=0, column=1)

        # Поля ввода
        self.entries = {}
        for i, label in enumerate(["Длина:", "Ширина:", "Сторона A:", "Сторона B:", "Сторона C:", "Высота:"]):
            tk.Label(self.root, text=label).grid(row=i + 1, column=0, sticky="w")
            entry = tk.Entry(self.root)
            entry.grid(row=i + 1, column=1)
            self.entries[label] = entry

        # Результаты
        tk.Label(self.root, text="Результаты:").grid(row=8, column=0, sticky="w")
        self.result_label = tk.Label(self.root, textvariable=self.result_text, fg="blue")
        self.result_label.grid(row=8, column=1)

        # Кнопки
        tk.Button(self.root, text="Рассчитать", command=self.calculate).grid(row=9, column=0)
        tk.Button(self.root, text="Сохранить", command=self.save_report).grid(row=9, column=1)

    def calculate(self):
        try:
            shape_type = self.shape_var.get()
            if shape_type == "Прямоугольник":
                length = float(self.entries["Длина:"].get())
                width = float(self.entries["Ширина:"].get())
                self.shape = Rectangle(length, width)

            elif shape_type == "Треугольник":
                a = float(self.entries["Сторона A:"].get())
                b = float(self.entries["Сторона B:"].get())
                c = float(self.entries["Сторона C:"].get())
                self.shape = Triangle(a, b, c)

            elif shape_type == "Трапеция":
                a = float(self.entries["Сторона A:"].get())
                b = float(self.entries["Сторона B:"].get())
                h = float(self.entries["Высота:"].get())
                self.shape = Trapezoid(a, b, h)

            area = self.shape.area()
            inscribed = self.shape.inscribed_radius()
            circumscribed = self.shape.circumscribed_radius()

            self.result_text.set(f"Площадь: {area:.2f}, Вписанный радиус: {inscribed}, Описанный радиус: {circumscribed}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка ввода данных: {e}")

    def save_report(self):
        try:
            filepath = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx"), ("Excel files", "*.xlsx")])
            if not filepath:
                return

            if filepath.endswith(".docx"):
                doc = Document()
                doc.add_heading("Результаты расчетов", level=1)
                doc.add_paragraph(self.result_text.get())
                doc.save(filepath)
            elif filepath.endswith(".xlsx"):
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "Результаты"
                ws.append(["Результаты", self.result_text.get()])
                wb.save(filepath)

            messagebox.showinfo("Успех", "Результаты успешно сохранены.")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка сохранения: {e}")

if __name__ == "__main__":
    root = tk.Tk()
    app = GeometryApp(root)
    root.mainloop()
